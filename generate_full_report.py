import os
import sys
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from barcode import EAN13
from barcode.writer import ImageWriter
import networkx as nx

# === Helper functions ===

def set_apa_format(paragraph, is_title=False, level=0):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if is_title:
            run.font.bold = True
            if level <= 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if level == 3:
                run.font.italic = True
    paragraph.paragraph_format.line_spacing = 2
    paragraph.paragraph_format.space_after = Pt(0)


def add_formatted(paragraph_text, doc, is_title=False, level=0):
    p = doc.add_paragraph(paragraph_text)
    set_apa_format(p, is_title=is_title, level=level)
    return p


# === Document generation ===

def generate_report(output_path='TRABAJO FINAL HPTA.docx'):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    # Portada simple
    add_formatted('REPORTE LOGÍSTICO', doc, is_title=True, level=0)
    doc.add_paragraph()
    add_formatted('Universidad INCCA de Colombia', doc)
    add_formatted('Curso: Logística Internacional', doc)
    doc.add_page_break()

    # Introducción
    add_formatted('Introducción', doc, is_title=True, level=1)
    intro = ('Este documento ejemplifica el uso de tablas, códigos de barras y diagramas '
             'para una propuesta logística internacional. Todas las secciones se presentan '
             'en formato APA 7 con letra Times New Roman a 12 puntos y doble espacio.')
    p = add_formatted(intro, doc)
    p.paragraph_format.first_line_indent = Inches(0.5)

    # === Matriz de códigos de barras e inventarios ===
    add_formatted('Matriz de códigos de barras e inventarios', doc, is_title=True, level=2)
    productos = [
        {'Producto': 'FURADAN 350ml', 'Codigo': '750103131111',
         'Inventario': 'EOQ', 'Vehiculo': 'Camion refrigerado',
         'Razon': 'Mantener temperatura'},
        {'Producto': 'Juxtapid 20mg', 'Codigo': '840123456789',
         'Inventario': 'ABC', 'Vehiculo': 'Camion seco',
         'Razon': 'Producto farmacéutico'},
        {'Producto': 'Alcohol Glicerinado 75%', 'Codigo': '770987654321',
         'Inventario': 'FIFO', 'Vehiculo': 'Contenedor estándar',
         'Razon': 'Alto volumen'}
    ]
    df_prod = pd.DataFrame(productos)

    table = doc.add_table(rows=1, cols=len(df_prod.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df_prod.columns):
        hdr_cells[i].text = col
    for _, row in df_prod.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    # Generar códigos de barras e insertar
    doc.add_paragraph()
    for prod in productos:
        code = EAN13(prod['Codigo'], writer=ImageWriter())
        filename = f"barcode_{prod['Codigo']}.png"
        code.save(filename)
        doc.add_picture(filename, width=Inches(1.5))
        os.remove(filename)

    # === Matriz del ciclo de compras ===
    add_formatted('Ciclo de compras', doc, is_title=True, level=2)
    etapas = ['Identificación de necesidad', 'Selección de proveedor',
              'Negociación', 'Pedido', 'Recepción', 'Pago']
    df_ciclo = pd.DataFrame({'Etapa': etapas, 'Descripción': ['']*len(etapas)})
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Etapa'
    table.rows[0].cells[1].text = 'Descripción'
    for _, row in df_ciclo.iterrows():
        r = table.add_row().cells
        r[0].text = row['Etapa']
        r[1].text = row['Descripción']

    # === Matriz de tipo de mercado y empresa ===
    add_formatted('Tipo de mercado y empresa proveedora', doc, is_title=True, level=2)
    mercado = pd.DataFrame({
        'Producto': ['FURADAN 350ml', 'Juxtapid 20mg', 'Alcohol Glicerinado 75%'],
        'Mercado': ['Internacional', 'Internacional', 'Nacional'],
        'Proveedor': ['Fabricante', 'Distribuidor', 'Fabricante']
    })
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    for i, col in enumerate(mercado.columns):
        table.rows[0].cells[i].text = col
    for _, row in mercado.iterrows():
        r = table.add_row().cells
        r[0].text = row['Producto']
        r[1].text = row['Mercado']
        r[2].text = row['Proveedor']

    # === Diagrama de rutas (simplificado) ===
    add_formatted('Diagrama de rutas', doc, is_title=True, level=2)
    G = nx.DiGraph()
    G.add_edge('Busan', 'Callao')
    G.add_edge('Pimpri-Chinchwad', 'Callao')
    G.add_edge('Callao', 'Juanjui')
    pos = nx.spring_layout(G)
    plt.figure()
    nx.draw(G, pos, with_labels=True, node_color='skyblue', node_size=1500, arrowsize=20)
    plt.savefig('rutas.png')
    plt.close()
    doc.add_picture('rutas.png', width=Inches(4))
    os.remove('rutas.png')

    # === Diagrama de cadena de suministro ===
    add_formatted('Cadena de suministro con IA', doc, is_title=True, level=2)
    H = nx.DiGraph()
    H.add_edge('Proveedores', 'Planta')
    H.add_edge('Planta', 'Centro de Distribución')
    H.add_edge('Centro de Distribución', 'Cliente')
    H.add_edge('Centro de Distribución', 'IA')
    pos = nx.spring_layout(H)
    plt.figure()
    nx.draw(H, pos, with_labels=True, node_color='lightgreen', node_size=1500, arrowsize=20)
    plt.savefig('supply_chain.png')
    plt.close()
    doc.add_picture('supply_chain.png', width=Inches(4))
    os.remove('supply_chain.png')

    # === Graficos adicionales ===
    add_formatted('Gráficos adicionales', doc, is_title=True, level=2)
    q = range(100, 1000, 100)
    costs = [0.5*x + 2000/x for x in q]
    plt.figure()
    plt.plot(q, costs)
    plt.xlabel('Cantidad')
    plt.ylabel('Costo total')
    plt.title('Costo total vs cantidad (EOQ)')
    plt.savefig('costos.png')
    plt.close()
    doc.add_picture('costos.png', width=Inches(4))
    os.remove('costos.png')

    doc.save(output_path)

if __name__ == '__main__':
    out_file = sys.argv[1] if len(sys.argv) > 1 else 'TRABAJO FINAL HPTA.docx'
    generate_report(out_file)
    print(f"Documento '{out_file}' creado correctamente.")
